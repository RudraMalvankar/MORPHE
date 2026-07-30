import io
import re
from abc import ABC, abstractmethod
from typing import List

import docx
import fitz  # PyMuPDF

from app.modules.cdm.schemas import (
    Author,
    CanonicalDocument,
    CitationRef,
    DocumentSection,
    MediaObject,
    SectionType,
)


class DocumentParser(ABC):
    @abstractmethod
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        """Parses raw file bytes into a clean CanonicalDocument CDM snapshot."""
        pass


class PDFParser(DocumentParser):
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        # Load PDF using PyMuPDF (fitz)
        pdf_stream = io.BytesIO(file_content)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")

        text = ""
        sections = []

        # Extract title from bookmarks or metadata
        title = doc.metadata.get("title") or "Untitled PDF Research Paper"
        author_name = doc.metadata.get("author") or "Unknown Author"

        # Split document by pages to extract text
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()

        # Simple segmenting strategy to identify sections/paragraphs (no NLP fallback)
        abstract = ""
        lines = text.split("\n")
        current_section_title = "Introduction"
        current_section_content: List[str] = []
        section_idx = 1

        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Detect Abstract
            if "abstract" in line.lower() and len(line) < 30:
                continue
            # Basic section header detection (capitalized line, short length)
            if line.isupper() and len(line) < 50 and not line.endswith("."):
                if current_section_content:
                    sections.append(
                        DocumentSection(
                            id=f"sec_{section_idx}",
                            section_type=SectionType.INTRODUCTION
                            if section_idx == 1
                            else SectionType.METHODOLOGY,
                            title=current_section_title,
                            content_markdown="\n".join(current_section_content),
                            order=section_idx,
                        )
                    )
                    section_idx += 1
                current_section_title = line.title()
                current_section_content = []
            else:
                if len(abstract) < 100 and "abstract" in line.lower():
                    abstract = line
                else:
                    current_section_content.append(line)

        # Append final section
        if current_section_content:
            sections.append(
                DocumentSection(
                    id=f"sec_{section_idx}",
                    section_type=SectionType.CONCLUSION,
                    title=current_section_title,
                    content_markdown="\n".join(current_section_content),
                    order=section_idx,
                )
            )

        # Basic References extraction
        references = []
        ref_sections = [
            s
            for s in sections
            if "reference" in s.title.lower() or "bibliography" in s.title.lower()
        ]
        if ref_sections:
            ref_content = ref_sections[0].content_markdown
            ref_lines = ref_content.split("\n")
            for idx, ref_line in enumerate(
                ref_lines[:20]
            ):  # Limit to 20 references for performance
                if len(ref_line) > 10:
                    references.append(
                        CitationRef(id=f"ref_{idx + 1}", raw_text=ref_line, title=ref_line[:50])
                    )

        # Detect images metadata (Part 5)
        media_objects = []
        media_idx = 1
        for page_num in range(len(doc)):
            images = doc.get_page_images(page_num)
            for img in images:
                media_objects.append(
                    MediaObject(
                        id=f"media_{media_idx}",
                        type="figure",
                        caption=f"Embedded Image {media_idx} on Page {page_num + 1}",
                        raw_content=f"width: {img[2]}, height: {img[3]}, bpc: {img[4]}",
                    )
                )
                media_idx += 1

        authors = [Author(id="auth_1", first_name=author_name, last_name="")]

        return CanonicalDocument(
            id=project_id,
            version_id=version_id,
            domain_profile_id="general",
            title=title,
            authors=authors,
            abstract=abstract or "Abstract not detected.",
            keywords=[],
            sections=sections,
            references=references,
            media_objects=media_objects,
            metadata={
                "page_count": len(doc),
                "producer": doc.metadata.get("producer", ""),
                "creator": doc.metadata.get("creator", ""),
            },
        )


class DocxParser(DocumentParser):
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        docx_stream = io.BytesIO(file_content)
        doc = docx.Document(docx_stream)

        title = "Untitled DOCX Document"
        abstract = ""
        sections = []
        references: List[CitationRef] = []
        media_objects = []

        current_section_title = "Header Info"
        current_section_content: List[str] = []
        section_idx = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Identify title
            if para.style is not None and getattr(para.style, "name", "") == "Title":
                title = text
                continue

            # Identify section headers
            if para.style is not None and getattr(para.style, "name", "").startswith("Heading"):
                if current_section_content:
                    sections.append(
                        DocumentSection(
                            id=f"sec_{section_idx}",
                            section_type=SectionType.INTRODUCTION
                            if section_idx == 1
                            else SectionType.METHODOLOGY,
                            title=current_section_title,
                            content_markdown="\n".join(current_section_content),
                            order=section_idx,
                        )
                    )
                    section_idx += 1
                current_section_title = text
                current_section_content = []
            else:
                if "abstract" in text.lower() and len(text) > 50:
                    abstract = text
                else:
                    current_section_content.append(text)

        # Append final section
        if current_section_content:
            sections.append(
                DocumentSection(
                    id=f"sec_{section_idx}",
                    section_type=SectionType.CONCLUSION,
                    title=current_section_title,
                    content_markdown="\n".join(current_section_content),
                    order=section_idx,
                )
            )

        # Image details parsing
        img_idx = 1
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                media_objects.append(
                    MediaObject(
                        id=f"media_{img_idx}",
                        type="figure",
                        caption=f"DOCX Inline Image {img_idx}",
                        raw_content=rel.target_ref,
                    )
                )
                img_idx += 1

        return CanonicalDocument(
            id=project_id,
            version_id=version_id,
            title=title,
            authors=[],
            abstract=abstract or "Abstract not detected.",
            keywords=[],
            sections=sections,
            references=references,
            media_objects=media_objects,
            metadata={"paragraph_count": len(doc.paragraphs)},
        )


class MarkdownParser(DocumentParser):
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        text = file_content.decode("utf-8", errors="ignore")
        lines = text.split("\n")

        title = "Untitled Markdown Document"
        abstract = ""
        sections = []
        references: List[CitationRef] = []
        media_objects: List[MediaObject] = []

        current_section_title = "Introduction"
        current_section_content: List[str] = []
        section_idx = 1

        # Check for YAML metadata blocks (Part 6)
        in_yaml = False
        yaml_lines: List[str] = []
        for line in lines:
            line_strip = line.strip()
            if line_strip == "---":
                if in_yaml:
                    in_yaml = False
                    # Parse basic title from YAML block
                    for y_line in yaml_lines:
                        if y_line.lower().startswith("title:"):
                            title = y_line.split(":", 1)[1].strip().strip('"').strip("'")
                else:
                    in_yaml = True
                continue
            if in_yaml:
                yaml_lines.append(line_strip)
                continue

            # Identify headers
            if line.startswith("#"):
                header_match = re.match(r"^(#+)\s+(.*)$", line_strip)
                if header_match:
                    level = len(header_match.group(1))
                    header_text = header_match.group(2)
                    if level == 1 and title == "Untitled Markdown Document":
                        title = header_text
                        continue

                    if current_section_content:
                        sections.append(
                            DocumentSection(
                                id=f"sec_{section_idx}",
                                section_type=SectionType.INTRODUCTION
                                if section_idx == 1
                                else SectionType.METHODOLOGY,
                                title=current_section_title,
                                content_markdown="\n".join(current_section_content),
                                order=section_idx,
                            )
                        )
                        section_idx += 1
                    current_section_title = header_text
                    current_section_content = []
            else:
                if "abstract" in line.lower() and len(line) > 50:
                    abstract = line
                else:
                    current_section_content.append(line)

        # Append final section
        if current_section_content:
            sections.append(
                DocumentSection(
                    id=f"sec_{section_idx}",
                    section_type=SectionType.CONCLUSION,
                    title=current_section_title,
                    content_markdown="\n".join(current_section_content),
                    order=section_idx,
                )
            )

        return CanonicalDocument(
            id=project_id,
            version_id=version_id,
            title=title,
            authors=[],
            abstract=abstract or "Abstract not detected.",
            keywords=[],
            sections=sections,
            references=references,
            media_objects=media_objects,
            metadata={},
        )


class LatexParser(DocumentParser):
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        text = file_content.decode("utf-8", errors="ignore")

        # Extract title using Regex (Part 7 LaTeX extraction)
        title_match = re.search(r"\\title\{([^}]+)\}", text)
        title = title_match.group(1) if title_match else "Untitled LaTeX Document"

        # Abstract
        abstract_match = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", text, re.DOTALL)
        abstract = abstract_match.group(1).strip() if abstract_match else "Abstract not detected."

        # Parse sections
        sections = []
        sec_matches = re.finditer(r"\\section\{([^}]+)\}(.*?)(?=\\section|\Z)", text, re.DOTALL)
        section_idx = 1
        for m in sec_matches:
            sec_title = m.group(1)
            sec_content = m.group(2).strip()
            # Clean basic latex command references
            clean_content = re.sub(r"\\[a-zA-Z]+(\[[^\]]*\])?\{([^}]+)\}", r"\2", sec_content)
            sections.append(
                DocumentSection(
                    id=f"sec_{section_idx}",
                    section_type=SectionType.INTRODUCTION
                    if section_idx == 1
                    else SectionType.METHODOLOGY,
                    title=sec_title,
                    content_markdown=clean_content,
                    order=section_idx,
                )
            )
            section_idx += 1

        # Fallback if no sections detected
        if not sections:
            sections.append(
                DocumentSection(
                    id="sec_1",
                    section_type=SectionType.INTRODUCTION,
                    title="Main Content",
                    content_markdown=text[:5000],
                    order=1,
                )
            )

        # Figure detection metadata (Part 7)
        media_objects = []
        fig_idx = 1
        fig_matches = re.finditer(r"\\begin\{figure\}(.*?)\\end\{figure\}", text, re.DOTALL)
        for m in fig_matches:
            caption_match = re.search(r"\\caption\{([^}]+)\}", m.group(1))
            caption = caption_match.group(1) if caption_match else f"LaTeX Figure {fig_idx}"
            media_objects.append(
                MediaObject(
                    id=f"media_{fig_idx}",
                    type="figure",
                    caption=caption,
                    raw_content=m.group(1).strip(),
                )
            )
            fig_idx += 1

        return CanonicalDocument(
            id=project_id,
            version_id=version_id,
            title=title,
            authors=[],
            abstract=abstract,
            keywords=[],
            sections=sections,
            references=[],
            media_objects=media_objects,
            metadata={},
        )


class PlainTextParser(DocumentParser):
    async def parse(
        self, file_content: bytes, project_id: str, version_id: str
    ) -> CanonicalDocument:
        text = file_content.decode("utf-8", errors="ignore")
        paragraphs = text.split("\n\n")

        title = paragraphs[0][:100] if paragraphs else "Untitled Plain Text Document"
        abstract = paragraphs[1] if len(paragraphs) > 1 else "Abstract not detected."

        sections = []
        section_idx = 1

        # Group paragraphs into a simple CDM structure
        for idx, p in enumerate(paragraphs[2:10]):
            sections.append(
                DocumentSection(
                    id=f"sec_{section_idx}",
                    section_type=SectionType.INTRODUCTION
                    if section_idx == 1
                    else SectionType.CONCLUSION,
                    title=f"Section {section_idx}",
                    content_markdown=p.strip(),
                    order=section_idx,
                )
            )
            section_idx += 1

        if not sections:
            sections.append(
                DocumentSection(
                    id="sec_1",
                    section_type=SectionType.INTRODUCTION,
                    title="Main Section",
                    content_markdown=text,
                    order=1,
                )
            )

        return CanonicalDocument(
            id=project_id,
            version_id=version_id,
            title=title,
            authors=[],
            abstract=abstract,
            keywords=[],
            sections=sections,
            references=[],
            media_objects=[],
            metadata={},
        )


class ParserFactory:
    @staticmethod
    def get_parser(mime_type: str) -> DocumentParser:
        mime = mime_type.lower()
        if "pdf" in mime:
            return PDFParser()
        elif "docx" in mime or "officedocument.wordprocessingml" in mime:
            return DocxParser()
        elif "markdown" in mime or "md" in mime:
            return MarkdownParser()
        elif "latex" in mime or "x-tex" in mime or "tex" in mime:
            return LatexParser()
        else:
            # Fallback to plain text
            return PlainTextParser()
