import uuid

import docx
import fitz
import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.core.deps import db_dependency
from app.db.base import Base
from app.main import app
from app.modules.auth.deps import create_access_token
from app.modules.auth.repository import UserRepository
from app.modules.ingestion.parsers import (
    DocxParser,
    LatexParser,
    MarkdownParser,
    PDFParser,
    PlainTextParser,
)
from app.modules.projects.repository import ProjectRepository

TEST_DATABASE_URL = "sqlite+aiosqlite:///:memory:"


@pytest_asyncio.fixture(scope="function")
async def db_session():
    engine = create_async_engine(TEST_DATABASE_URL, echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    async_session_local = async_sessionmaker(
        bind=engine,
        class_=AsyncSession,
        expire_on_commit=False,
    )

    async with async_session_local() as session:
        yield session
        await session.rollback()

    await engine.dispose()


@pytest_asyncio.fixture(scope="function")
async def client(db_session):
    async def override_db():
        yield db_session

    from app.modules.ingestion.router import get_session_maker

    class MockSessionMaker:
        def __init__(self, session):
            self.session = session

        def __call__(self):
            return self

        async def __aenter__(self):
            return self.session

        async def __aexit__(self, exc_type, exc_val, exc_tb):
            pass

    app.dependency_overrides[db_dependency] = override_db
    app.dependency_overrides[get_session_maker] = lambda: MockSessionMaker(db_session)
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        yield ac
    app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_pdf_parser():
    # Build dummy PDF bytes
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "ABSTRACT\nThis is the dummy abstract content.\n"
        "INTRODUCTION\nMain content of introduction.",
    )
    pdf_bytes = doc.write()
    doc.close()

    parser = PDFParser()
    cdm = await parser.parse(pdf_bytes, project_id=str(uuid.uuid4()), version_id=str(uuid.uuid4()))
    assert cdm.title is not None
    assert "introduction" in cdm.sections[0].title.lower()


@pytest.mark.asyncio
async def test_docx_parser():
    # Build dummy docx bytes
    doc = docx.Document()
    doc.add_heading("Quantum Core Layouts", level=0)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("First paragraph text details.")
    import io

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    parser = DocxParser()
    cdm = await parser.parse(docx_bytes, project_id=str(uuid.uuid4()), version_id=str(uuid.uuid4()))
    assert cdm.title == "Quantum Core Layouts"
    assert cdm.sections[0].title == "Introduction"


@pytest.mark.asyncio
async def test_markdown_parser():
    md_content = """---
title: "Markdown Research Paper"
---
# Markdown Research Paper
## Introduction
This is md text.
"""
    parser = MarkdownParser()
    cdm = await parser.parse(
        md_content.encode("utf-8"), project_id=str(uuid.uuid4()), version_id=str(uuid.uuid4())
    )
    assert cdm.title == "Markdown Research Paper"
    assert cdm.sections[0].title == "Introduction"


@pytest.mark.asyncio
async def test_latex_parser():
    latex_content = """
\\title{LaTeX Core Platform}
\\begin{abstract}
This abstract details LaTeX formats.
\\end{abstract}
\\section{Introduction}
Text details here.
"""
    parser = LatexParser()
    cdm = await parser.parse(
        latex_content.encode("utf-8"), project_id=str(uuid.uuid4()), version_id=str(uuid.uuid4())
    )
    assert cdm.title == "LaTeX Core Platform"
    assert cdm.sections[0].title == "Introduction"


@pytest.mark.asyncio
async def test_plain_text_parser():
    txt_content = "Plain Text Research Paper\n\nAbstract details\n\nIntroduction paragraph content."
    parser = PlainTextParser()
    cdm = await parser.parse(
        txt_content.encode("utf-8"), project_id=str(uuid.uuid4()), version_id=str(uuid.uuid4())
    )
    assert cdm.title == "Plain Text Research Paper"
    assert "Introduction" in cdm.sections[0].content_markdown


@pytest.mark.asyncio
async def test_ingestion_api_endpoints(client, db_session):
    user_repo = UserRepository(db_session)
    user = await user_repo.create_user("user@morphe.edu", "pass123", "Alice")

    proj_repo = ProjectRepository(db_session)
    project = await proj_repo.create(user.id, "Quantum Analysis")

    token = create_access_token(str(user.id))
    headers = {"Authorization": f"Bearer {token}"}

    # Upload and trigger async parse
    file_payload = {
        "file": (
            "paper.txt",
            b"Title: Research Paper\n\nAbstract: details\n\nIntroduction text content.",
            "text/plain",
        )
    }
    response = await client.post(
        f"/api/v1/ingestion/upload?project_id={project.id}", headers=headers, files=file_payload
    )
    assert response.status_code == 202
    job_id = response.json()["job_id"]

    # Check job status
    job_res = await client.get(f"/api/v1/ingestion/jobs/{job_id}", headers=headers)
    assert job_res.status_code == 200
    assert job_res.json()["status"] in {"queued", "running", "completed"}
